from django.contrib.auth.decorators import login_required
from .models import PatientProfile
from .forms import PatientProfileForm
from core.decorators import group_required
from django.views.decorators.csrf import csrf_exempt
from doctor.models import DoctorProfile
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse
from django.utils import timezone
from .models import Appointment, MedicalRecord
import json
import datetime
from django.contrib.auth.models import User
from django.utils import timezone
from datetime import timedelta, datetime
from django.utils.timezone import make_aware
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
from django.templatetags.static import static
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
import os
from pymongo import MongoClient
import gridfs
from bson import ObjectId
from patient.management.commands.runtelegrambot import schedule_appointment_notifications
from datetime import datetime
from asgiref.sync import async_to_sync
from patient.management.commands.runtelegrambot import send_notification


@login_required
@group_required('patient')
def patient_myrecords(request):
    appointments = Appointment.objects.filter(patient=request.user)
    return render(request, 'patient/myrecords.html', {'appointments': appointments})

@login_required
@group_required('patient')
def delete_appointment(request, appointment_id):
    appointment = get_object_or_404(Appointment, id=appointment_id, patient=request.user)
    if request.method == 'POST':
        appointment_info = f'{appointment.date} в {appointment.time} к доктору {appointment.doctor.first_name} {appointment.doctor.last_name}'
        try:
            chat_id = appointment.patient.telegramuser.chat_id
            async_to_sync(send_notification)(chat_id, f'Ваша запись {appointment_info} отменена.')
        except User.telegramuser.RelatedObjectDoesNotExist:
            pass
        appointment.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False}, status=400)

@csrf_exempt
@login_required
@group_required('patient')
def new_appointment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        doctor_id = data.get('doctor_id')
        date_str = data.get('date')
        time_str = data.get('time')

        date = datetime.strptime(date_str, "%Y-%m-%d").date()
        time = datetime.strptime(time_str, "%H:%M").time()

        doctor = DoctorProfile.objects.get(id=doctor_id)

        appointment = Appointment.objects.create(
            patient=request.user,
            doctor=doctor,
            date=date,
            time=time
        )

        try:
            schedule_appointment_notifications(appointment)
        except User.telegramuser.RelatedObjectDoesNotExist:
            pass

        return JsonResponse({'success': True})
    return JsonResponse({'success': False})

@login_required
@group_required('patient')
def get_doctors(request):
    doctors = DoctorProfile.objects.values('id', 'first_name', 'last_name', 'patronymic', 'specialization', 'hospital_address')
    return JsonResponse(list(doctors), safe=False)

@login_required
@group_required('patient')
def get_available_times(request):
    today = timezone.now().date()
    available_times = []

    # Логика генерации доступных дат и времен
    i = 0
    while len(available_times) < 2:
        day = today + timedelta(days=i)
        if day.weekday() < 5:  # Будние дни
            times = []
            for hour in [9, 11, 13, 15, 17]:
                naive_datetime = datetime.combine(day, datetime.min.time()) + timedelta(hours=hour)
                aware_datetime = make_aware(naive_datetime, timezone.get_current_timezone())
                if aware_datetime > timezone.now():
                    times.append(aware_datetime.strftime('%H:%M'))
            if times:
                available_times.append({
                    'date': day.strftime('%Y-%m-%d'),
                    'times': times
                })
        i += 1

    return JsonResponse(available_times, safe=False)

@login_required
@group_required('patient')
def patient_myinfo(request):
    user = request.user
    profile, created = PatientProfile.objects.get_or_create(user=user)  # Получаем или создаем профиль
    if request.method == 'POST':
        form = PatientProfileForm(request.POST, instance=profile)
        if form.is_valid():
            form.save()
            return redirect('patient:myinfo')
    else:
        form = PatientProfileForm(instance=profile)
    groups = user.groups.all()
    return render(request, 'patient/myinfo.html', {'user': user, 'groups': groups, 'form': form, 'profile': profile})

@login_required
@group_required('patient')
def patient_mymedicalcard(request):
    user = request.user
    medical_records = MedicalRecord.objects.filter(patient=user).order_by('-date_completed')
    context = {
        'patient': user,
        'medical_records': medical_records,
    }
    return render(request, 'patient/mymedicalcard.html', context)

@login_required
@group_required('patient')
def export_medical_record_pdf(request, record_id):
    record = get_object_or_404(MedicalRecord, id=record_id)
    client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
    db = client[settings.MONGO_DB['db']]
    fs = gridfs.GridFS(db)
    
    if record.image_id:
        image = fs.get(ObjectId(record.image_id))
        image_data = image.read()
        image_url = f'/tmp/{image.filename}'
        with open(image_url, 'wb') as f:
            f.write(image_data)
        image_url = request.build_absolute_uri(f'/media/{image.filename}')
    else:
        image_url = None
    
    html_string = render_to_string('patient/medical_record_pdf.html', {'record': record, 'image_url': image_url})
    html = HTML(string=html_string, base_url=request.build_absolute_uri('/'))
    
    css_url = static('core/styles.css')
    css_path = os.path.join(settings.STATIC_ROOT, css_url.replace(settings.STATIC_URL, ''))
    css_files = [
        CSS(string='@page { size: A4; margin: 1cm }'),
        CSS(css_path),
    ]
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="medical_record_{record_id}.pdf"'
    html.write_pdf(response, stylesheets=css_files)
    return response

@login_required
@group_required('patient')
def export_medical_record_doc(request, record_id):
    record = get_object_or_404(MedicalRecord, id=record_id)
    document = Document()
    document.add_heading('Медицинская запись', 0)

    document.add_heading('Описание приёма', level=1)
    document.add_paragraph(record.description)

    document.add_heading('Заключение', level=1)
    document.add_paragraph(record.conclusion)

    document.add_heading('Дата завершения', level=1)
    document.add_paragraph(str(record.date_completed))

    if record.image_id:
        client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
        db = client[settings.MONGO_DB['db']]
        fs = gridfs.GridFS(db)
        image = fs.get(ObjectId(record.image_id))
        image_data = image.read()
        image_path = f'/tmp/{image.filename}'
        with open(image_path, 'wb') as f:
            f.write(image_data)
        document.add_heading('Изображение', level=1)
        document.add_picture(image_path, width=Inches(4.0))
        os.remove(image_path)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="medical_record_{record_id}.docx"'
    document.save(response)
    return response

@login_required
@group_required('patient')
def load_image(request, image_id):
    client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
    db = client[settings.MONGO_DB['db']]
    fs = gridfs.GridFS(db)
    image_id = ObjectId(image_id)
    image = fs.get(image_id)
    response = HttpResponse(image.read(), content_type='image/jpeg')
    response['Content-Disposition'] = f'inline; filename={image.filename}'
    return response