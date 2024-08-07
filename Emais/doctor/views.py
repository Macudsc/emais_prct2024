from .forms import DoctorProfileForm
from django.contrib.auth.models import User
import logging
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from core.decorators import group_required
from django.http import JsonResponse
from .models import DoctorProfile
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.templatetags.static import static
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
import os
from django.http import HttpResponse, JsonResponse
from django.urls import reverse
from django.utils.html import format_html
from pymongo import MongoClient
import gridfs
from bson import ObjectId
from django.views.decorators.http import require_http_methods
from patient.models import MedicalRecord, Appointment
from asgiref.sync import async_to_sync
from patient.management.commands.runtelegrambot import send_notification
from django.core.files.uploadedfile import UploadedFile
import tempfile
import re

#Моя информация
@login_required
@group_required('doctor')
def doctor_myinfo(request):
    user = request.user
    profile, created = DoctorProfile.objects.get_or_create(user=user)  # Получаем или создаем профиль
    if request.method == 'POST':
        form = DoctorProfileForm(request.POST, instance=profile)
        if form.is_valid():
            form.save()
            return redirect('doctor:myinfo')
    else:
        form = DoctorProfileForm(instance=profile)
    groups = user.groups.all()
    return render(request, 'doctor/myinfo.html', {'user': user, 'groups': groups, 'form': form, 'profile': profile})

# Приём пациентов
@login_required
@group_required('doctor')
def doctor_mypatients(request):
    doctor = DoctorProfile.objects.get(user=request.user)
    appointments = Appointment.objects.filter(doctor=doctor)
    context = {
        'doctor': doctor,
        'appointments': appointments,
    }
    return render(request, 'doctor/mypatients.html', context)

# Заполнить исследование
@login_required
@group_required('doctor')
def complete_appointment(request, appointment_id):
    appointment = get_object_or_404(Appointment, id=appointment_id)
    if request.method == 'POST':
        description = request.POST.get('description')
        conclusion = request.POST.get('conclusion')
        image = request.FILES.get('upload')

        max_size_mb = 3
        if image and isinstance(image, UploadedFile) and image.size > max_size_mb * 1024 * 1024:
            return JsonResponse({'status': 'failed', 'error': 'Размер файла превышает 3 МБ.'})

        medical_record = MedicalRecord(
            patient=appointment.patient,
            doctor=appointment.doctor,
            appointment=appointment,
            description=description,
            conclusion=conclusion
        )
        if image:
            medical_record.save_image(image)
        else:
            medical_record.save()

        try:
            chat_id = appointment.patient.telegramuser.chat_id
            async_to_sync(send_notification)(chat_id, f'Ваш приём завершён и исследование доступно к просмотру. Ссылка на медкарту: https://5gr4vbqg-8000.euw.devtunnels.ms/patient/mymedicalcard/')
        except User.telegramuser.RelatedObjectDoesNotExist:
            pass

        appointment.delete() # ! ИГРУШКА ДЬЯВОЛА
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'failed', 'error': 'Invalid request method'})

# Просмотр медкарты
@login_required
@group_required('doctor')
@require_http_methods(["GET"])
def view_medical_records(request, patient_id):
    medical_records = MedicalRecord.objects.filter(patient_id=patient_id).order_by('-date_completed')
    return render(request, 'doctor/medical_records.html', {'medical_records': medical_records})

@login_required
@group_required('doctor')
def export_medical_record_pdf(request, record_id):
    record = get_object_or_404(MedicalRecord, id=record_id)
    client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
    db = client[settings.MONGO_DB['db']]
    fs = gridfs.GridFS(db)

    image_url = None
    if record.image_id:
        image = fs.get(ObjectId(record.image_id))
        image_data = image.read()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
            temp_file.write(image_data)
            temp_file.flush()
            image_url = temp_file.name

    html_string = render_to_string('patient/medical_record_pdf.html', {'record': record, 'image_url': image_url})
    html = HTML(string=html_string, base_url=request.build_absolute_uri('/'))

    css_url = static('core/styles.css')
    css_path = os.path.join(settings.STATIC_ROOT, css_url.replace(settings.STATIC_URL, ''))
    css_files = [
        CSS(string='@page { size: A4; margin: 1cm }'),
        CSS(css_path),
    ]

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="medical_record_{record_id}.pdf"'
    html.write_pdf(response, stylesheets=css_files)

    if image_url and os.path.exists(image_url):
        os.remove(image_url)

    return response


@login_required
@group_required('doctor')
def export_medical_record_doc(request, record_id):
    record = get_object_or_404(MedicalRecord, id=record_id)
    document = Document()
    document.add_heading('Медицинская запись', 0)

    def clean_text(text):
        text = text.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    document.add_heading('Описание приёма', level=1)
    clean_description = clean_text(record.description)
    document.add_paragraph(clean_description, style='BodyText')

    document.add_heading('Заключение', level=1)
    clean_conclusion = clean_text(record.conclusion)
    document.add_paragraph(clean_conclusion, style='BodyText')

    document.add_heading('Дата завершения', level=1)
    document.add_paragraph(str(record.date_completed), style='BodyText')

    if record.image_id:
        client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
        db = client[settings.MONGO_DB['db']]
        fs = gridfs.GridFS(db)
        image = fs.get(ObjectId(record.image_id))
        image_data = image.read()
        image_path = f'/tmp/{image.filename}'
        with open(image_path, 'wb') as f:
            f.write(image_data)
        document.add_heading('Изображение', level=1)
        document.add_picture(image_path, width=Inches(4.0))
        os.remove(image_path)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="medical_record_{record_id}.docx"'
    document.save(response)
    return response

@login_required
@group_required('doctor')
def load_image(request, image_id):
    client = MongoClient(settings.MONGO_DB['host'], settings.MONGO_DB['port'])
    db = client[settings.MONGO_DB['db']]
    fs = gridfs.GridFS(db)
    image_id = ObjectId(image_id)
    image = fs.get(image_id)
    response = HttpResponse(image.read(), content_type='image/jpeg')
    response['Content-Disposition'] = f'inline; filename={image.filename}'
    return response
